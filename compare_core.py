import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from xlsx_exporter import save_compare_xlsx

from settings_service import APP_NAME, APP_SUBTITLE, APP_VERSION, PRODUCT_DOMAIN, PUBLISHER_NAME, SUITE_NAME, TOOL_FOLDER_NAME, ensure_directories


ALGORITHM_KEYS = {
    "MD5": "md5",
    "SHA-1": "sha1",
    "SHA-256": "sha256"
}


def safe_filename(value, fallback="compare"):
    value = (value or "").strip()

    if not value:
        value = fallback

    invalid_chars = '<>:"/\\|?*'

    for char in invalid_chars:
        value = value.replace(char, "_")

    value = value.replace(" ", "_")

    while "__" in value:
        value = value.replace("__", "_")

    return value.strip("_") or fallback


def load_manifest_json(path):
    manifest_path = Path(path)

    with manifest_path.open("r", encoding="utf-8") as f:
        manifest = json.load(f)

    if not isinstance(manifest, dict):
        raise ValueError("Manifest JSON did not contain an object.")

    if "files" not in manifest or not isinstance(manifest["files"], list):
        raise ValueError("Manifest JSON does not contain a valid files list.")

    return manifest


def get_manifest_algorithms(manifest):
    hash_settings = manifest.get("hash_settings", {})
    algorithms = hash_settings.get("algorithms", [])

    if not isinstance(algorithms, list):
        return []

    cleaned = []

    for algorithm in algorithms:
        algorithm = str(algorithm).strip().upper()

        if algorithm == "SHA1":
            algorithm = "SHA-1"
        elif algorithm == "SHA256":
            algorithm = "SHA-256"

        if algorithm in ALGORITHM_KEYS and algorithm not in cleaned:
            cleaned.append(algorithm)

    return cleaned


def get_common_algorithms(manifest_a, manifest_b):
    algorithms_a = get_manifest_algorithms(manifest_a)
    algorithms_b = get_manifest_algorithms(manifest_b)

    common = []

    for algorithm in ["SHA-256", "SHA-1", "MD5"]:
        if algorithm in algorithms_a and algorithm in algorithms_b:
            common.append(algorithm)

    if common:
        return common

    # Fallback for older or incomplete manifests.
    return ["SHA-256", "SHA-1", "MD5"]


def get_match_key(record):
    relative_path = str(record.get("relative_path", "")).strip()

    if relative_path:
        return relative_path.lower()

    file_name = str(record.get("file_name", "")).strip()

    if file_name:
        return file_name.lower()

    file_path = str(record.get("file_path", "")).strip()

    return file_path.lower()


def build_index(records):
    index = {}

    for record in records:
        key = get_match_key(record)

        if not key:
            continue

        index.setdefault(key, []).append(record)

    return index


def get_hash(record, algorithm_name):
    key = ALGORITHM_KEYS.get(algorithm_name)

    if not key:
        return None

    value = record.get(key)

    if value is None:
        return None

    value = str(value).strip()

    if not value or value.lower() == "not calculated":
        return None

    return value.lower()


def compare_file_records(record_a, record_b, algorithms):
    hashes_a = {}
    hashes_b = {}
    compared = []

    for algorithm in algorithms:
        value_a = get_hash(record_a, algorithm)
        value_b = get_hash(record_b, algorithm)

        hashes_a[algorithm] = value_a or "Not calculated"
        hashes_b[algorithm] = value_b or "Not calculated"

        if value_a and value_b:
            compared.append((algorithm, value_a, value_b))

    if not compared:
        return {
            "status": "Error",
            "detail": "No comparable hash values were available for this file.",
            "algorithms_checked": [],
            "manifest_a_hashes": hashes_a,
            "manifest_b_hashes": hashes_b
        }

    mismatches = []

    for algorithm, value_a, value_b in compared:
        if value_a != value_b:
            mismatches.append(algorithm)

    if mismatches:
        return {
            "status": "Hash mismatch",
            "detail": "Hash mismatch detected for: " + ", ".join(mismatches),
            "algorithms_checked": [item[0] for item in compared],
            "manifest_a_hashes": hashes_a,
            "manifest_b_hashes": hashes_b
        }

    return {
        "status": "Matched",
        "detail": "Compared hash values matched.",
        "algorithms_checked": [item[0] for item in compared],
        "manifest_a_hashes": hashes_a,
        "manifest_b_hashes": hashes_b
    }


def build_compare_report(manifest_a, manifest_b, manifest_a_path, manifest_b_path, technician, notes):
    files_a = manifest_a.get("files", [])
    files_b = manifest_b.get("files", [])

    algorithms = get_common_algorithms(manifest_a, manifest_b)

    index_a = build_index(files_a)
    index_b = build_index(files_b)

    results = []
    matched_b_keys = set()

    for key, records_a in index_a.items():
        records_b = index_b.get(key)

        if len(records_a) > 1:
            for record_a in records_a:
                results.append({
                    "file_name": record_a.get("file_name", ""),
                    "match_key": key,
                    "status": "Error",
                    "detail": "Duplicate file key found in Manifest A. Comparison is ambiguous.",
                    "manifest_a_path": record_a.get("file_path", ""),
                    "manifest_b_path": "",
                    "algorithms_checked": [],
                    "manifest_a_hashes": {},
                    "manifest_b_hashes": {}
                })
            continue

        record_a = records_a[0]

        if not records_b:
            results.append({
                "file_name": record_a.get("file_name", ""),
                "match_key": key,
                "status": "Only in Manifest A",
                "detail": "File exists in Manifest A but not in Manifest B.",
                "manifest_a_path": record_a.get("file_path", ""),
                "manifest_b_path": "",
                "algorithms_checked": [],
                "manifest_a_hashes": {},
                "manifest_b_hashes": {}
            })
            continue

        if len(records_b) > 1:
            matched_b_keys.add(key)
            results.append({
                "file_name": record_a.get("file_name", ""),
                "match_key": key,
                "status": "Error",
                "detail": "Duplicate file key found in Manifest B. Comparison is ambiguous.",
                "manifest_a_path": record_a.get("file_path", ""),
                "manifest_b_path": "",
                "algorithms_checked": [],
                "manifest_a_hashes": {},
                "manifest_b_hashes": {}
            })
            continue

        record_b = records_b[0]
        matched_b_keys.add(key)

        if record_a.get("hash_status") == "Error" or record_b.get("hash_status") == "Error":
            results.append({
                "file_name": record_a.get("file_name", "") or record_b.get("file_name", ""),
                "match_key": key,
                "status": "Error",
                "detail": "At least one manifest contains an error status for this file.",
                "manifest_a_path": record_a.get("file_path", ""),
                "manifest_b_path": record_b.get("file_path", ""),
                "algorithms_checked": [],
                "manifest_a_hashes": {},
                "manifest_b_hashes": {}
            })
            continue

        comparison = compare_file_records(record_a, record_b, algorithms)

        results.append({
            "file_name": record_a.get("file_name", "") or record_b.get("file_name", ""),
            "match_key": key,
            "status": comparison["status"],
            "detail": comparison["detail"],
            "manifest_a_path": record_a.get("file_path", ""),
            "manifest_b_path": record_b.get("file_path", ""),
            "algorithms_checked": comparison["algorithms_checked"],
            "manifest_a_hashes": comparison["manifest_a_hashes"],
            "manifest_b_hashes": comparison["manifest_b_hashes"]
        })

    for key, records_b in index_b.items():
        if key in matched_b_keys:
            continue

        for record_b in records_b:
            results.append({
                "file_name": record_b.get("file_name", ""),
                "match_key": key,
                "status": "Only in Manifest B",
                "detail": "File exists in Manifest B but not in Manifest A.",
                "manifest_a_path": "",
                "manifest_b_path": record_b.get("file_path", ""),
                "algorithms_checked": [],
                "manifest_a_hashes": {},
                "manifest_b_hashes": {}
            })

    summary = summarize_compare_results(results)

    report = {
        "app_name": APP_NAME,
        "app_subtitle": APP_SUBTITLE,
        "app_version": APP_VERSION,
        "suite_name": SUITE_NAME,
        "publisher": PUBLISHER_NAME,
        "product_domain": PRODUCT_DOMAIN,
        "tool_folder_name": TOOL_FOLDER_NAME,
        "compare_created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "compare_type": "Manifest compare",
        "technician": technician.strip(),
        "manifest_a": {
            "path": manifest_a_path,
            "app_name": manifest_a.get("app_name", ""),
            "app_version": manifest_a.get("app_version", ""),
            "created_at": manifest_a.get("created_at", ""),
            "case_info": manifest_a.get("case_info", {})
        },
        "manifest_b": {
            "path": manifest_b_path,
            "app_name": manifest_b.get("app_name", ""),
            "app_version": manifest_b.get("app_version", ""),
            "created_at": manifest_b.get("created_at", ""),
            "case_info": manifest_b.get("case_info", {})
        },
        "algorithms_considered": algorithms,
        "summary": summary,
        "results": results,
        "notes": notes.strip()
    }

    return report


def summarize_compare_results(results):
    summary = {
        "total_results": len(results),
        "matched": 0,
        "hash_mismatch": 0,
        "only_in_manifest_a": 0,
        "only_in_manifest_b": 0,
        "errors": 0
    }

    for result in results:
        status = result.get("status", "")

        if status == "Matched":
            summary["matched"] += 1
        elif status == "Hash mismatch":
            summary["hash_mismatch"] += 1
        elif status == "Only in Manifest A":
            summary["only_in_manifest_a"] += 1
        elif status == "Only in Manifest B":
            summary["only_in_manifest_b"] += 1
        elif status == "Error":
            summary["errors"] += 1

    return summary


def build_compare_txt(report):
    manifest_a = report.get("manifest_a", {})
    manifest_b = report.get("manifest_b", {})
    case_a = manifest_a.get("case_info", {})
    case_b = manifest_b.get("case_info", {})
    summary = report.get("summary", {})
    results = report.get("results", [])

    lines = []

    lines.append("BYTECASE VERIFY")
    lines.append("Hash Manifest Compare Report")
    lines.append("=" * 80)
    lines.append(f"Part of the {report.get('suite_name', SUITE_NAME)} toolset by {report.get('publisher', PUBLISHER_NAME)}")
    lines.append(f"Product Domain: {report.get('product_domain', PRODUCT_DOMAIN)}")
    lines.append("")

    lines.append("COMPARE INFORMATION")
    lines.append("-" * 80)
    lines.append(f"Generated By: {report.get('app_name', '')} - {report.get('app_subtitle', '')} v{report.get('app_version', '')}")
    lines.append(f"Compare Created At: {report.get('compare_created_at', '')}")
    lines.append(f"Technician: {report.get('technician', '')}")
    lines.append(f"Algorithms Considered: {', '.join(report.get('algorithms_considered', []))}")
    lines.append("")

    lines.append("MANIFEST A")
    lines.append("-" * 80)
    lines.append(f"Path: {manifest_a.get('path', '')}")
    lines.append(f"Original App: {manifest_a.get('app_name', '')} v{manifest_a.get('app_version', '')}")
    lines.append(f"Created At: {manifest_a.get('created_at', '')}")
    lines.append(f"Case Number: {case_a.get('case_number', '')}")
    lines.append(f"Agency Case Number: {case_a.get('agency_case_number', '')}")
    lines.append(f"Technician: {case_a.get('technician', '')}")
    lines.append(f"Source Description: {case_a.get('source_description', '')}")
    lines.append("")

    lines.append("MANIFEST B")
    lines.append("-" * 80)
    lines.append(f"Path: {manifest_b.get('path', '')}")
    lines.append(f"Original App: {manifest_b.get('app_name', '')} v{manifest_b.get('app_version', '')}")
    lines.append(f"Created At: {manifest_b.get('created_at', '')}")
    lines.append(f"Case Number: {case_b.get('case_number', '')}")
    lines.append(f"Agency Case Number: {case_b.get('agency_case_number', '')}")
    lines.append(f"Technician: {case_b.get('technician', '')}")
    lines.append(f"Source Description: {case_b.get('source_description', '')}")
    lines.append("")

    lines.append("COMPARE SUMMARY")
    lines.append("-" * 80)
    lines.append(f"Total Results: {summary.get('total_results', 0)}")
    lines.append(f"Matched: {summary.get('matched', 0)}")
    lines.append(f"Hash Mismatch: {summary.get('hash_mismatch', 0)}")
    lines.append(f"Only in Manifest A: {summary.get('only_in_manifest_a', 0)}")
    lines.append(f"Only in Manifest B: {summary.get('only_in_manifest_b', 0)}")
    lines.append(f"Errors: {summary.get('errors', 0)}")
    lines.append("")

    notes = report.get("notes", "")

    if notes:
        lines.append("NOTES")
        lines.append("-" * 80)
        lines.append(notes)
        lines.append("")

    lines.append("COMPARE RESULTS")
    lines.append("-" * 80)
    lines.append("")

    for index, result in enumerate(results, start=1):
        lines.append(f"Result {index}")
        lines.append("-" * 40)
        lines.append(f"File Name: {result.get('file_name', '')}")
        lines.append(f"Status: {result.get('status', '')}")
        lines.append(f"Detail: {result.get('detail', '')}")
        lines.append(f"Manifest A Path: {result.get('manifest_a_path', '')}")
        lines.append(f"Manifest B Path: {result.get('manifest_b_path', '')}")
        lines.append(f"Algorithms Checked: {', '.join(result.get('algorithms_checked', []))}")

        hashes_a = result.get("manifest_a_hashes", {})
        hashes_b = result.get("manifest_b_hashes", {})

        if hashes_a or hashes_b:
            lines.append("Manifest A Hashes:")

            for algorithm, value in hashes_a.items():
                lines.append(f"  {algorithm}: {value}")

            lines.append("Manifest B Hashes:")

            for algorithm, value in hashes_b.items():
                lines.append(f"  {algorithm}: {value}")

        lines.append("")

    lines.append("=" * 80)
    lines.append("End of ByteCase Verify Hash Manifest Compare Report")
    lines.append("=" * 80)

    return "\n".join(lines)


def save_compare_csv(report, csv_path):
    results = report.get("results", [])

    fieldnames = [
        "file_name",
        "status",
        "detail",
        "manifest_a_path",
        "manifest_b_path",
        "algorithms_checked",
        "manifest_a_md5",
        "manifest_b_md5",
        "manifest_a_sha1",
        "manifest_b_sha1",
        "manifest_a_sha256",
        "manifest_b_sha256"
    ]

    with csv_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()

        for result in results:
            hashes_a = result.get("manifest_a_hashes", {})
            hashes_b = result.get("manifest_b_hashes", {})

            writer.writerow({
                "file_name": result.get("file_name", ""),
                "status": result.get("status", ""),
                "detail": result.get("detail", ""),
                "manifest_a_path": result.get("manifest_a_path", ""),
                "manifest_b_path": result.get("manifest_b_path", ""),
                "algorithms_checked": ", ".join(result.get("algorithms_checked", [])),
                "manifest_a_md5": hashes_a.get("MD5", ""),
                "manifest_b_md5": hashes_b.get("MD5", ""),
                "manifest_a_sha1": hashes_a.get("SHA-1", ""),
                "manifest_b_sha1": hashes_b.get("SHA-1", ""),
                "manifest_a_sha256": hashes_a.get("SHA-256", ""),
                "manifest_b_sha256": hashes_b.get("SHA-256", "")
            })


def save_compare_docx(report, docx_path):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_font = document.styles["Normal"].font
    normal_font.name = "Arial"
    normal_font.size = Pt(10)

    document.add_heading("Hash Manifest Compare Report", level=0)

    manifest_a = report.get("manifest_a", {})
    manifest_b = report.get("manifest_b", {})
    case_a = manifest_a.get("case_info", {})
    case_b = manifest_b.get("case_info", {})
    summary = report.get("summary", {})

    document.add_heading("Compare Information", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Generated By", f"{report.get('app_name', '')} v{report.get('app_version', '')}"),
        ("Compare Created At", report.get("compare_created_at", "")),
        ("Technician", report.get("technician", "")),
        ("Algorithms Considered", ", ".join(report.get("algorithms_considered", [])))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Manifest A", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Path", manifest_a.get("path", "")),
        ("Original App", f"{manifest_a.get('app_name', '')} v{manifest_a.get('app_version', '')}"),
        ("Created At", manifest_a.get("created_at", "")),
        ("Case Number", case_a.get("case_number", "")),
        ("Agency Case Number", case_a.get("agency_case_number", "")),
        ("Technician", case_a.get("technician", "")),
        ("Source Description", case_a.get("source_description", ""))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Manifest B", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Path", manifest_b.get("path", "")),
        ("Original App", f"{manifest_b.get('app_name', '')} v{manifest_b.get('app_version', '')}"),
        ("Created At", manifest_b.get("created_at", "")),
        ("Case Number", case_b.get("case_number", "")),
        ("Agency Case Number", case_b.get("agency_case_number", "")),
        ("Technician", case_b.get("technician", "")),
        ("Source Description", case_b.get("source_description", ""))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Compare Summary", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Total Results", summary.get("total_results", 0)),
        ("Matched", summary.get("matched", 0)),
        ("Hash Mismatch", summary.get("hash_mismatch", 0)),
        ("Only in Manifest A", summary.get("only_in_manifest_a", 0)),
        ("Only in Manifest B", summary.get("only_in_manifest_b", 0)),
        ("Errors", summary.get("errors", 0))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    notes = report.get("notes", "")

    if notes:
        document.add_heading("Notes", level=1)
        document.add_paragraph(notes)

    document.add_heading("Compare Results", level=1)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = [
        "#",
        "File Name",
        "Status",
        "Detail",
        "Manifest A Path",
        "Manifest B Path"
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for index, result in enumerate(report.get("results", []), start=1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = result.get("file_name", "")
        cells[2].text = result.get("status", "")
        cells[3].text = result.get("detail", "")
        cells[4].text = result.get("manifest_a_path", "")
        cells[5].text = result.get("manifest_b_path", "")

    document.add_paragraph()
    document.add_paragraph("End of Hash Manifest Compare Report")

    document.save(docx_path)


def save_compare_outputs(report, settings):
    manifest_a = report.get("manifest_a", {})
    case_info_for_paths = manifest_a.get("case_info", {})
    paths = ensure_directories(settings, case_number=case_info_for_paths.get("case_number", ""), mode_folder="comparisons")
    case_info = manifest_a.get("case_info", {})

    case_number = case_info.get("case_number", "")
    source_description = case_info.get("source_description", "")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    safe_case = safe_filename(case_number, "NO_CASE")
    safe_source = safe_filename(source_description, "compare")
    base_filename = f"{safe_case}_{safe_source}_{timestamp}_compare_report"

    txt_path = paths["reports_dir"] / f"{base_filename}.txt"
    csv_path = paths["reports_dir"] / f"{base_filename}.csv"
    docx_path = paths["reports_dir"] / f"{base_filename}.docx"
    xlsx_path = paths["reports_dir"] / f"{base_filename}.xlsx"
    json_path = paths["saved_manifests_dir"] / f"{base_filename}.json"

    with txt_path.open("w", encoding="utf-8") as f:
        f.write(build_compare_txt(report))

    save_compare_csv(report, csv_path)
    save_compare_docx(report, docx_path)
    save_compare_xlsx(report, xlsx_path)

    with json_path.open("w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)

    return txt_path, csv_path, docx_path, xlsx_path, json_path