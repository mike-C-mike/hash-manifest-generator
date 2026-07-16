import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from xlsx_exporter import save_verification_xlsx

from settings_service import APP_NAME, APP_SUBTITLE, APP_VERSION, PRODUCT_DOMAIN, PUBLISHER_NAME, SUITE_NAME, TOOL_FOLDER_NAME, ensure_directories


ALGORITHM_KEYS = {
    "MD5": "md5",
    "SHA-1": "sha1",
    "SHA-256": "sha256"
}


def safe_filename(value, fallback="verification"):
    value = (value or "").strip()

    if not value:
        value = fallback

    invalid_chars = '<>:"/\\|?*'

    for char in invalid_chars:
        value = value.replace(char, "_")

    value = value.replace(" ", "_")

    while "__" in value:
        value = value.replace("__", "_")

    return value.strip("_") or fallback


def format_bytes(size_bytes):
    if size_bytes is None:
        return "Unknown"

    size = float(size_bytes)

    for unit in ["bytes", "KB", "MB", "GB", "TB"]:
        if size < 1024 or unit == "TB":
            if unit == "bytes":
                return f"{int(size)} {unit}"

            return f"{size:.2f} {unit}"

        size /= 1024

    return f"{size_bytes} bytes"


def load_manifest_json(path):
    """
    Loads a previously generated manifest JSON file.
    """
    manifest_path = Path(path)

    with manifest_path.open("r", encoding="utf-8") as f:
        manifest = json.load(f)

    if not isinstance(manifest, dict):
        raise ValueError("Manifest JSON did not contain an object.")

    if "files" not in manifest or not isinstance(manifest["files"], list):
        raise ValueError("Manifest JSON does not contain a valid files list.")

    return manifest


def get_manifest_algorithms(manifest):
    """
    Returns algorithms listed in a prior manifest.
    """
    hash_settings = manifest.get("hash_settings", {})
    algorithms = hash_settings.get("algorithms", [])

    if not isinstance(algorithms, list):
        return []

    cleaned = []

    for algorithm in algorithms:
        algorithm = str(algorithm).strip().upper()

        if algorithm == "SHA1":
            algorithm = "SHA-1"
        elif algorithm == "SHA256":
            algorithm = "SHA-256"

        if algorithm in ALGORITHM_KEYS and algorithm not in cleaned:
            cleaned.append(algorithm)

    return cleaned


def get_match_key(record):
    """
    Returns the comparison key for a file record.

    Current v0.x manifests usually store relative_path as the filename.
    Later versions can improve this with true folder-relative paths.
    """
    relative_path = str(record.get("relative_path", "")).strip()

    if relative_path:
        return relative_path.lower()

    file_name = str(record.get("file_name", "")).strip()

    if file_name:
        return file_name.lower()

    file_path = str(record.get("file_path", "")).strip()

    return file_path.lower()


def build_index(records):
    """
    Builds a dictionary index for comparison.

    Duplicate keys are retained as lists so the report can flag ambiguity.
    """
    index = {}

    for record in records:
        key = get_match_key(record)

        if not key:
            continue

        index.setdefault(key, []).append(record)

    return index


def get_hash(record, algorithm_name):
    key = ALGORITHM_KEYS.get(algorithm_name)

    if not key:
        return None

    value = record.get(key)

    if value is None:
        return None

    value = str(value).strip()

    if not value or value.lower() == "not calculated":
        return None

    return value.lower()


def compare_records(original_record, current_record, algorithms):
    original_hashes = {}
    current_hashes = {}
    compared = []

    for algorithm in algorithms:
        original_value = get_hash(original_record, algorithm)
        current_value = get_hash(current_record, algorithm)

        original_hashes[algorithm] = original_value or "Not calculated"
        current_hashes[algorithm] = current_value or "Not calculated"

        if original_value and current_value:
            compared.append((algorithm, original_value, current_value))

    if not compared:
        return {
            "status": "Error",
            "detail": "No comparable hash values were available for this file.",
            "algorithms_checked": [],
            "original_hashes": original_hashes,
            "current_hashes": current_hashes
        }

    mismatches = []

    for algorithm, original_value, current_value in compared:
        if original_value != current_value:
            mismatches.append(algorithm)

    if mismatches:
        return {
            "status": "Hash mismatch",
            "detail": "Hash mismatch detected for: " + ", ".join(mismatches),
            "algorithms_checked": [item[0] for item in compared],
            "original_hashes": original_hashes,
            "current_hashes": current_hashes
        }

    return {
        "status": "Matched",
        "detail": "Compared hash values matched.",
        "algorithms_checked": [item[0] for item in compared],
        "original_hashes": original_hashes,
        "current_hashes": current_hashes
    }


def build_verification_report(original_manifest, current_files, technician, notes):
    """
    Compares current hash results against a previously generated manifest.
    """
    original_files = original_manifest.get("files", [])
    original_case_info = original_manifest.get("case_info", {})
    algorithms = get_manifest_algorithms(original_manifest)

    if not algorithms:
        algorithms = ["SHA-256", "SHA-1", "MD5"]

    original_index = build_index(original_files)
    current_index = build_index(current_files)

    results = []
    matched_current_keys = set()

    for key, original_records in original_index.items():
        current_records = current_index.get(key)

        if len(original_records) > 1:
            for original_record in original_records:
                results.append({
                    "file_name": original_record.get("file_name", ""),
                    "match_key": key,
                    "status": "Error",
                    "detail": "Duplicate file key found in original manifest. Verification is ambiguous.",
                    "original_path": original_record.get("file_path", ""),
                    "current_path": "",
                    "algorithms_checked": [],
                    "original_hashes": {},
                    "current_hashes": {}
                })
            continue

        original_record = original_records[0]

        if not current_records:
            results.append({
                "file_name": original_record.get("file_name", ""),
                "match_key": key,
                "status": "Missing from current selection",
                "detail": "File from original manifest was not found in current selection.",
                "original_path": original_record.get("file_path", ""),
                "current_path": "",
                "algorithms_checked": [],
                "original_hashes": {},
                "current_hashes": {}
            })
            continue

        if len(current_records) > 1:
            matched_current_keys.add(key)
            results.append({
                "file_name": original_record.get("file_name", ""),
                "match_key": key,
                "status": "Error",
                "detail": "Duplicate file key found in current selection. Verification is ambiguous.",
                "original_path": original_record.get("file_path", ""),
                "current_path": "",
                "algorithms_checked": [],
                "original_hashes": {},
                "current_hashes": {}
            })
            continue

        current_record = current_records[0]
        matched_current_keys.add(key)

        if current_record.get("hash_status") == "Error":
            results.append({
                "file_name": original_record.get("file_name", ""),
                "match_key": key,
                "status": "Error",
                "detail": current_record.get("error", "Current file could not be hashed."),
                "original_path": original_record.get("file_path", ""),
                "current_path": current_record.get("file_path", ""),
                "algorithms_checked": [],
                "original_hashes": {},
                "current_hashes": {}
            })
            continue

        comparison = compare_records(original_record, current_record, algorithms)

        results.append({
            "file_name": original_record.get("file_name", ""),
            "match_key": key,
            "status": comparison["status"],
            "detail": comparison["detail"],
            "original_path": original_record.get("file_path", ""),
            "current_path": current_record.get("file_path", ""),
            "algorithms_checked": comparison["algorithms_checked"],
            "original_hashes": comparison["original_hashes"],
            "current_hashes": comparison["current_hashes"]
        })

    for key, current_records in current_index.items():
        if key in matched_current_keys:
            continue

        for current_record in current_records:
            results.append({
                "file_name": current_record.get("file_name", ""),
                "match_key": key,
                "status": "New file not in original manifest",
                "detail": "Current file was not listed in the original manifest.",
                "original_path": "",
                "current_path": current_record.get("file_path", ""),
                "algorithms_checked": [],
                "original_hashes": {},
                "current_hashes": {}
            })

    summary = summarize_verification_results(results)

    report = {
        "app_name": APP_NAME,
        "app_subtitle": APP_SUBTITLE,
        "app_version": APP_VERSION,
        "suite_name": SUITE_NAME,
        "publisher": PUBLISHER_NAME,
        "product_domain": PRODUCT_DOMAIN,
        "tool_folder_name": TOOL_FOLDER_NAME,
        "verification_created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "verification_type": "Manifest verification",
        "technician": technician.strip(),
        "original_manifest": {
            "app_name": original_manifest.get("app_name", ""),
            "app_version": original_manifest.get("app_version", ""),
            "created_at": original_manifest.get("created_at", ""),
            "case_info": original_case_info
        },
        "algorithms_considered": algorithms,
        "summary": summary,
        "results": results,
        "notes": notes.strip()
    }

    return report


def summarize_verification_results(results):
    summary = {
        "total_results": len(results),
        "matched": 0,
        "hash_mismatch": 0,
        "missing_from_current_selection": 0,
        "new_file_not_in_original_manifest": 0,
        "errors": 0
    }

    for result in results:
        status = result.get("status", "")

        if status == "Matched":
            summary["matched"] += 1
        elif status == "Hash mismatch":
            summary["hash_mismatch"] += 1
        elif status == "Missing from current selection":
            summary["missing_from_current_selection"] += 1
        elif status == "New file not in original manifest":
            summary["new_file_not_in_original_manifest"] += 1
        elif status == "Error":
            summary["errors"] += 1

    return summary


def build_verification_txt(report):
    original = report.get("original_manifest", {})
    case_info = original.get("case_info", {})
    summary = report.get("summary", {})
    results = report.get("results", [])

    lines = []

    lines.append("BYTECASE VERIFY")
    lines.append("Hash Manifest Verification Report")
    lines.append("=" * 80)
    lines.append(f"Part of the {report.get('suite_name', SUITE_NAME)} toolset by {report.get('publisher', PUBLISHER_NAME)}")
    lines.append(f"Product Domain: {report.get('product_domain', PRODUCT_DOMAIN)}")
    lines.append("")

    lines.append("VERIFICATION INFORMATION")
    lines.append("-" * 80)
    lines.append(f"Generated By: {report.get('app_name', '')} - {report.get('app_subtitle', '')} v{report.get('app_version', '')}")
    lines.append(f"Verification Created At: {report.get('verification_created_at', '')}")
    lines.append(f"Technician: {report.get('technician', '')}")
    lines.append(f"Algorithms Considered: {', '.join(report.get('algorithms_considered', []))}")
    lines.append("")

    lines.append("ORIGINAL MANIFEST INFORMATION")
    lines.append("-" * 80)
    lines.append(f"Original App: {original.get('app_name', '')} v{original.get('app_version', '')}")
    lines.append(f"Original Created At: {original.get('created_at', '')}")
    lines.append(f"Case Number: {case_info.get('case_number', '')}")
    lines.append(f"Agency Case Number: {case_info.get('agency_case_number', '')}")
    lines.append(f"Original Technician: {case_info.get('technician', '')}")
    lines.append(f"Source Description: {case_info.get('source_description', '')}")
    lines.append("")

    lines.append("VERIFICATION SUMMARY")
    lines.append("-" * 80)
    lines.append(f"Total Results: {summary.get('total_results', 0)}")
    lines.append(f"Matched: {summary.get('matched', 0)}")
    lines.append(f"Hash Mismatch: {summary.get('hash_mismatch', 0)}")
    lines.append(f"Missing From Current Selection: {summary.get('missing_from_current_selection', 0)}")
    lines.append(f"New File Not In Original Manifest: {summary.get('new_file_not_in_original_manifest', 0)}")
    lines.append(f"Errors: {summary.get('errors', 0)}")
    lines.append("")

    notes = report.get("notes", "")

    if notes:
        lines.append("NOTES")
        lines.append("-" * 80)
        lines.append(notes)
        lines.append("")

    lines.append("VERIFICATION RESULTS")
    lines.append("-" * 80)
    lines.append("")

    for index, result in enumerate(results, start=1):
        lines.append(f"Result {index}")
        lines.append("-" * 40)
        lines.append(f"File Name: {result.get('file_name', '')}")
        lines.append(f"Status: {result.get('status', '')}")
        lines.append(f"Detail: {result.get('detail', '')}")
        lines.append(f"Original Path: {result.get('original_path', '')}")
        lines.append(f"Current Path: {result.get('current_path', '')}")
        lines.append(f"Algorithms Checked: {', '.join(result.get('algorithms_checked', []))}")

        original_hashes = result.get("original_hashes", {})
        current_hashes = result.get("current_hashes", {})

        if original_hashes or current_hashes:
            lines.append("Original Hashes:")

            for algorithm, value in original_hashes.items():
                lines.append(f"  {algorithm}: {value}")

            lines.append("Current Hashes:")

            for algorithm, value in current_hashes.items():
                lines.append(f"  {algorithm}: {value}")

        lines.append("")

    lines.append("=" * 80)
    lines.append("End of ByteCase Verify Hash Manifest Verification Report")
    lines.append("=" * 80)

    return "\n".join(lines)


def save_verification_csv(report, csv_path):
    results = report.get("results", [])

    fieldnames = [
        "file_name",
        "status",
        "detail",
        "original_path",
        "current_path",
        "algorithms_checked",
        "original_md5",
        "current_md5",
        "original_sha1",
        "current_sha1",
        "original_sha256",
        "current_sha256"
    ]

    with csv_path.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()

        for result in results:
            original_hashes = result.get("original_hashes", {})
            current_hashes = result.get("current_hashes", {})

            writer.writerow({
                "file_name": result.get("file_name", ""),
                "status": result.get("status", ""),
                "detail": result.get("detail", ""),
                "original_path": result.get("original_path", ""),
                "current_path": result.get("current_path", ""),
                "algorithms_checked": ", ".join(result.get("algorithms_checked", [])),
                "original_md5": original_hashes.get("MD5", ""),
                "current_md5": current_hashes.get("MD5", ""),
                "original_sha1": original_hashes.get("SHA-1", ""),
                "current_sha1": current_hashes.get("SHA-1", ""),
                "original_sha256": original_hashes.get("SHA-256", ""),
                "current_sha256": current_hashes.get("SHA-256", "")
            })


def save_verification_docx(report, docx_path):
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal_font = document.styles["Normal"].font
    normal_font.name = "Arial"
    normal_font.size = Pt(10)

    document.add_heading("Hash Manifest Verification Report", level=0)

    original = report.get("original_manifest", {})
    case_info = original.get("case_info", {})
    summary = report.get("summary", {})

    document.add_heading("Verification Information", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Generated By", f"{report.get('app_name', '')} v{report.get('app_version', '')}"),
        ("Verification Created At", report.get("verification_created_at", "")),
        ("Technician", report.get("technician", "")),
        ("Algorithms Considered", ", ".join(report.get("algorithms_considered", [])))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Original Manifest Information", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Original App", f"{original.get('app_name', '')} v{original.get('app_version', '')}"),
        ("Original Created At", original.get("created_at", "")),
        ("Case Number", case_info.get("case_number", "")),
        ("Agency Case Number", case_info.get("agency_case_number", "")),
        ("Original Technician", case_info.get("technician", "")),
        ("Source Description", case_info.get("source_description", ""))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    document.add_heading("Verification Summary", level=1)
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Total Results", summary.get("total_results", 0)),
        ("Matched", summary.get("matched", 0)),
        ("Hash Mismatch", summary.get("hash_mismatch", 0)),
        ("Missing From Current Selection", summary.get("missing_from_current_selection", 0)),
        ("New File Not In Original Manifest", summary.get("new_file_not_in_original_manifest", 0)),
        ("Errors", summary.get("errors", 0))
    ]

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)

    notes = report.get("notes", "")

    if notes:
        document.add_heading("Notes", level=1)
        document.add_paragraph(notes)

    document.add_heading("Verification Results", level=1)

    table = document.add_table(rows=1, cols=6)
    table.style = "Table Grid"

    headers = [
        "#",
        "File Name",
        "Status",
        "Detail",
        "Original Path",
        "Current Path"
    ]

    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for index, result in enumerate(report.get("results", []), start=1):
        cells = table.add_row().cells
        cells[0].text = str(index)
        cells[1].text = result.get("file_name", "")
        cells[2].text = result.get("status", "")
        cells[3].text = result.get("detail", "")
        cells[4].text = result.get("original_path", "")
        cells[5].text = result.get("current_path", "")

    document.add_paragraph()
    document.add_paragraph("End of Hash Manifest Verification Report")

    document.save(docx_path)


def save_verification_outputs(report, settings):
    original = report.get("original_manifest", {})
    case_info_for_paths = original.get("case_info", {})
    paths = ensure_directories(settings, case_number=case_info_for_paths.get("case_number", ""), mode_folder="verifications")
    case_info = original.get("case_info", {})

    case_number = case_info.get("case_number", "")
    source_description = case_info.get("source_description", "")
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    safe_case = safe_filename(case_number, "NO_CASE")
    safe_source = safe_filename(source_description, "verification")
    base_filename = f"{safe_case}_{safe_source}_{timestamp}_verification_report"

    txt_path = paths["reports_dir"] / f"{base_filename}.txt"
    csv_path = paths["reports_dir"] / f"{base_filename}.csv"
    docx_path = paths["reports_dir"] / f"{base_filename}.docx"
    xlsx_path = paths["reports_dir"] / f"{base_filename}.xlsx"
    json_path = paths["saved_manifests_dir"] / f"{base_filename}.json"

    with txt_path.open("w", encoding="utf-8") as f:
        f.write(build_verification_txt(report))

    save_verification_csv(report, csv_path)
    save_verification_docx(report, docx_path)
    save_verification_xlsx(report, xlsx_path)

    with json_path.open("w", encoding="utf-8") as f:
        json.dump(report, f, indent=2)

    return txt_path, csv_path, docx_path, xlsx_path, json_path