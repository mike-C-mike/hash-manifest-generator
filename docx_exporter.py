from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from report_templates import HASH_GENERATION_METHOD, HASHING_EXPLANATION


SUPPORTED_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}


def format_bytes(size_bytes):
    """
    Human-readable file size.
    """
    if size_bytes is None:
        return "Unknown"

    size = float(size_bytes)

    for unit in ["bytes", "KB", "MB", "GB", "TB"]:
        if size < 1024 or unit == "TB":
            if unit == "bytes":
                return f"{int(size)} {unit}"

            return f"{size:.2f} {unit}"

        size /= 1024

    return f"{size_bytes} bytes"


def set_document_defaults(document):
    """
    Applies basic document formatting.
    """
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles

    normal_style = styles["Normal"]
    normal_font = normal_style.font
    normal_font.name = "Arial"
    normal_font.size = Pt(10)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def add_branding_image(document, settings):
    """
    Adds optional patch/logo image at the top of the report.

    Missing, invalid, or unsupported images are silently skipped so report generation
    does not fail because of branding.
    """
    report_branding = settings.get("report_branding", {})
    patch_image_path = report_branding.get("patch_image_path", "").strip()

    if not patch_image_path:
        return

    image_path = Path(patch_image_path)

    if not image_path.exists() or not image_path.is_file():
        return

    if image_path.suffix.lower() not in SUPPORTED_IMAGE_EXTENSIONS:
        return

    try:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(1.25))
    except Exception:
        # Branding should never block report generation.
        return


def add_key_value_table(document, rows):
    """
    Adds a simple two-column key/value table.
    """
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"

    for label, value in rows:
        row_cells = table.add_row().cells
        row_cells[0].text = str(label)
        row_cells[1].text = str(value)

    document.add_paragraph()


def add_file_summary_table(document, summary):
    """
    Adds file summary table.
    """
    rows = [
        ("Total Files Listed", summary.get("total_files", 0)),
        ("Completed", summary.get("completed_count", 0)),
        ("Errors", summary.get("error_count", 0)),
        ("Total Size", format_bytes(summary.get("total_size_bytes", 0)))
    ]

    add_key_value_table(document, rows)


def add_hashed_files_table(document, files):
    """
    Adds a table with one row per hashed file.
    """
    table = document.add_table(rows=1, cols=8)
    table.style = "Table Grid"

    headers = [
        "#",
        "File Name",
        "Size",
        "Status",
        "MD5",
        "SHA-1",
        "SHA-256",
        "Path"
    ]

    header_cells = table.rows[0].cells

    for index, header in enumerate(headers):
        header_cells[index].text = header

    for idx, record in enumerate(files, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(record.get("file_name", ""))
        row_cells[2].text = format_bytes(record.get("file_size_bytes"))
        row_cells[3].text = str(record.get("hash_status", ""))
        row_cells[4].text = record.get("md5") or "Not calculated"
        row_cells[5].text = record.get("sha1") or "Not calculated"
        row_cells[6].text = record.get("sha256") or "Not calculated"
        row_cells[7].text = str(record.get("file_path", ""))

    document.add_paragraph()


def add_error_details(document, files):
    """
    Adds error details if any files failed to hash.
    """
    error_records = [record for record in files if record.get("hash_status") == "Error"]

    if not error_records:
        return

    document.add_heading("Hashing Errors", level=2)

    for index, record in enumerate(error_records, start=1):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"Error {index}: ").bold = True
        paragraph.add_run(record.get("file_name", ""))

        document.add_paragraph(f"Path: {record.get('file_path', '')}")
        document.add_paragraph(f"Error: {record.get('error', '')}")


def add_signature_block(document, manifest):
    case_info = manifest.get("case_info", {})

    document.add_heading("Report Review / Signature", level=1)

    add_key_value_table(
        document,
        [
            ("Technician", case_info.get("technician", "")),
            ("Reviewed By", case_info.get("reviewed_by", "")),
            ("Date", datetime.now().strftime("%Y-%m-%d"))
        ]
    )

    document.add_paragraph()
    document.add_paragraph("Technician Signature: _______________________________")
    document.add_paragraph()
    document.add_paragraph("Reviewer Signature: _______________________________")


def build_docx_manifest(manifest, settings):
    """
    Builds the DOCX hash manifest report.
    """
    document = Document()
    set_document_defaults(document)

    add_branding_image(document, settings)

    department = manifest.get("department", {})
    case_info = manifest.get("case_info", {})
    hash_settings = manifest.get("hash_settings", {})
    report_options = manifest.get("report_options", {})
    file_summary = manifest.get("file_summary", {})
    files = manifest.get("files", [])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ByteCase Verify - Hash Manifest")
    title_run.bold = True
    title_run.font.size = Pt(16)

    department_name = department.get("department_name", "")
    unit_name = department.get("unit_name", "")

    if department_name or unit_name:
        dept_paragraph = document.add_paragraph()
        dept_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if department_name:
            dept_paragraph.add_run(department_name)

        if department_name and unit_name:
            dept_paragraph.add_run("\n")

        if unit_name:
            dept_paragraph.add_run(unit_name)

    document.add_paragraph()

    document.add_heading("Manifest Information", level=1)
    add_key_value_table(
        document,
        [
            ("Generated By", f"{manifest.get('app_name', '')} - {manifest.get('app_subtitle', '')} v{manifest.get('app_version', '')}"),
            ("Suite", manifest.get("suite_name", "ByteCase")),
            ("Publisher", manifest.get("publisher", "Forensics Byte")),
            ("Product Domain", manifest.get("product_domain", "byte-case.com")),
            ("Created At", manifest.get("created_at", ""))
        ]
    )

    document.add_heading("Case Information", level=1)
    add_key_value_table(
        document,
        [
            ("Case Number", case_info.get("case_number", "")),
            ("Agency Case Number", case_info.get("agency_case_number", "")),
            ("Exhibit / Item Reference", case_info.get("exhibit_reference", "")),
            ("Technician", case_info.get("technician", "")),
            ("Reviewed By", case_info.get("reviewed_by", "")),
            ("Source Description", case_info.get("source_description", ""))
        ]
    )

    document.add_heading("Hash Settings", level=1)
    algorithms = hash_settings.get("algorithms", [])
    add_key_value_table(
        document,
        [
            ("Algorithms", ", ".join(algorithms) if algorithms else "None selected"),
            ("Recursive Folder Selection", "Yes" if hash_settings.get("recursive") else "No"),
            ("Hash Generation Method Included", "Yes" if hash_settings.get("include_hash_generation_method") else "No"),
            ("Hashing Explanation Included", "Yes" if hash_settings.get("include_hashing_explanation") else "No"),
            ("Signature Block Included", "Yes" if report_options.get("include_signature_block") else "No")
        ]
    )

    if hash_settings.get("include_hash_generation_method"):
        document.add_heading("Hash Generation Method", level=1)
        document.add_paragraph(HASH_GENERATION_METHOD)

    document.add_heading("File Summary", level=1)
    add_file_summary_table(document, file_summary)

    if hash_settings.get("include_hashing_explanation"):
        document.add_heading("Hashing Explanation", level=1)
        document.add_paragraph(HASHING_EXPLANATION)

    notes = manifest.get("notes", "")

    if notes:
        document.add_heading("Notes", level=1)
        document.add_paragraph(notes)

    document.add_heading("Hashed Files", level=1)
    add_hashed_files_table(document, files)

    add_error_details(document, files)

    if report_options.get("include_signature_block"):
        add_signature_block(document, manifest)

    document.add_paragraph()
    end_paragraph = document.add_paragraph("End of ByteCase Verify Hash Manifest")
    end_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return document


def save_docx_manifest(manifest, settings, docx_path):
    """
    Saves the DOCX manifest report.
    """
    document = build_docx_manifest(manifest, settings)
    document.save(docx_path)
    return docx_path